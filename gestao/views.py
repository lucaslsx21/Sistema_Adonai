from django.db.models import Count
from django.db.models.functions import TruncMonth
from django.db import transaction

from decimal import Decimal
from datetime import datetime, date
import json
import openpyxl

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.contrib import messages

from openpyxl import Workbook
from docx import Document
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader

from .models import (
    CategoriaProduto,
    Doador,
    Produto,
    Campanha,
    Doacao,
    Beneficiario,
    Coleta,
    MovimentacaoEstoque,
    EntregaBeneficiario,
    DocumentoImportado,
    RegistroImportado,
)

from .forms import (
    DoadorForm,
    DocumentoImportadoForm,
    ProdutoForm,
    CampanhaForm,
    DoacaoForm,
    BeneficiarioForm,
    ColetaForm,
    MovimentacaoEstoqueForm,
    EntregaBeneficiarioForm,
)


# =========================
# FUNÇÕES AUXILIARES
# =========================

def normalizar_texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def normalizar_cabecalho(valor):
    return (
        normalizar_texto(valor)
        .lower()
        .replace(" ", "_")
        .replace("ç", "c")
        .replace("ã", "a")
        .replace("á", "a")
        .replace("à", "a")
        .replace("â", "a")
        .replace("é", "e")
        .replace("ê", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ô", "o")
        .replace("õ", "o")
        .replace("ú", "u")
    )


def montar_linha_dict(cabecalhos, linha):
    dados = {}
    for indice, cabecalho in enumerate(cabecalhos):
        if cabecalho:
            dados[cabecalho] = linha[indice] if indice < len(linha) else None
    return dados


def valor_decimal(valor, padrao=0):
    try:
        if valor in [None, ""]:
            return Decimal(padrao)
        return Decimal(str(valor).replace(",", "."))
    except Exception:
        return Decimal(padrao)


def valor_int(valor, padrao=0):
    try:
        if valor in [None, ""]:
            return padrao
        return int(float(valor))
    except Exception:
        return padrao


def valor_bool(valor):
    texto = normalizar_texto(valor).lower()
    return texto in ["sim", "s", "ativo", "true", "1", "yes"]


def valor_data(valor):
    if not valor:
        return None

    if isinstance(valor, datetime):
        return valor.date()

    if isinstance(valor, date):
        return valor

    texto = normalizar_texto(valor)

    for formato in ["%d/%m/%Y", "%Y-%m-%d"]:
        try:
            return datetime.strptime(texto, formato).date()
        except Exception:
            pass

    return None


def abrir_excel_com_dados(caminho):
    workbook = openpyxl.load_workbook(caminho, data_only=True)
    sheet = workbook.active

    linhas = list(sheet.iter_rows(values_only=True))

    if not linhas:
        return []

    cabecalhos = [normalizar_cabecalho(celula) for celula in linhas[0]]
    dados = []

    for linha in linhas[1:]:
        if not any(linha):
            continue
        dados.append(montar_linha_dict(cabecalhos, linha))

    return dados


# =========================
# IMPORTAÇÃO EXCEL
# =========================

def importar_produtos_excel(caminho):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            nome = normalizar_texto(
                item.get("nome") or item.get("produto") or item.get("descricao")
            )

            if not nome:
                continue

            categoria_nome = normalizar_texto(item.get("categoria") or "Sem categoria")
            unidade = normalizar_texto(item.get("unidade") or "UN").upper()
            estoque_minimo = valor_decimal(item.get("estoque_minimo") or item.get("minimo"), 0)
            quantidade = valor_decimal(item.get("quantidade") or item.get("estoque") or item.get("qtd"), 0)

            categoria, _ = CategoriaProduto.objects.get_or_create(nome=categoria_nome)

            produto, criado = Produto.objects.get_or_create(
                nome=nome,
                defaults={
                    "categoria": categoria,
                    "unidade": unidade,
                    "estoque_minimo": estoque_minimo,
                }
            )

            if not criado:
                produto.categoria = categoria
                produto.unidade = unidade
                produto.estoque_minimo = estoque_minimo
                produto.save()

            if quantidade > 0:
                MovimentacaoEstoque.objects.create(
                    produto=produto,
                    tipo="ENTRADA",
                    quantidade=quantidade,
                    observacao="Entrada automática por importação de documento."
                )

            total += 1

    return total


def importar_doadores_excel(caminho):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            nome = normalizar_texto(item.get("nome") or item.get("doador"))

            if not nome:
                continue

            tipo = normalizar_texto(item.get("tipo") or "PF").upper()

            if tipo not in ["PF", "PJ"]:
                tipo = "PF"

            Doador.objects.update_or_create(
                nome=nome,
                defaults={
                    "tipo": tipo,
                    "telefone": normalizar_texto(item.get("telefone")),
                    "email": normalizar_texto(item.get("email")),
                    "cidade": normalizar_texto(item.get("cidade")),
                    "endereco": normalizar_texto(item.get("endereco")),
                    "ativo": valor_bool(item.get("ativo") or "sim"),
                }
            )

            total += 1

    return total


def importar_beneficiarios_excel(caminho):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            nome = normalizar_texto(item.get("nome") or item.get("beneficiario"))

            if not nome:
                continue

            Beneficiario.objects.update_or_create(
                nome=nome,
                defaults={
                    "cpf": normalizar_texto(item.get("cpf")),
                    "telefone": normalizar_texto(item.get("telefone")),
                    "endereco": normalizar_texto(item.get("endereco")),
                    "quantidade_pessoas": valor_int(
                        item.get("quantidade_pessoas") or item.get("pessoas"),
                        1
                    ),
                    "observacoes": normalizar_texto(
                        item.get("observacoes") or item.get("observacao")
                    ),
                    "ativo": valor_bool(item.get("ativo") or "sim"),
                }
            )

            total += 1

    return total


def importar_campanhas_excel(caminho):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            nome = normalizar_texto(item.get("nome") or item.get("campanha"))

            if not nome:
                continue

            Campanha.objects.update_or_create(
                nome=nome,
                defaults={
                    "descricao": normalizar_texto(item.get("descricao")),
                    "data_inicio": valor_data(item.get("data_inicio") or item.get("inicio")) or date.today(),
                    "data_fim": valor_data(item.get("data_fim") or item.get("fim")) or date.today(),
                    "ativa": valor_bool(item.get("ativa") or "sim"),
                }
            )

            total += 1

    return total


def importar_estoque_excel(caminho):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            produto_nome = normalizar_texto(item.get("produto") or item.get("nome"))

            if not produto_nome:
                continue

            produto = Produto.objects.filter(nome__iexact=produto_nome).first()

            if not produto:
                categoria, _ = CategoriaProduto.objects.get_or_create(nome="Sem categoria")

                produto = Produto.objects.create(
                    nome=produto_nome,
                    categoria=categoria,
                    unidade="UN",
                    estoque_minimo=0,
                )

            tipo = normalizar_texto(item.get("tipo") or "ENTRADA").upper()

            if tipo not in ["ENTRADA", "SAIDA"]:
                tipo = "ENTRADA"

            quantidade = valor_decimal(item.get("quantidade") or item.get("qtd"), 0)

            if quantidade <= 0:
                continue

            MovimentacaoEstoque.objects.create(
                produto=produto,
                tipo=tipo,
                quantidade=quantidade,
                observacao=normalizar_texto(
                    item.get("observacao") or "Movimentação automática por importação."
                )
            )

            total += 1

    return total


def importar_doacoes_excel(caminho, usuario):
    linhas = abrir_excel_com_dados(caminho)
    total = 0

    with transaction.atomic():
        for item in linhas:
            nome_doador = normalizar_texto(item.get("doador") or item.get("nome_doador"))
            nome_campanha = normalizar_texto(item.get("campanha"))
            data_recebimento = valor_data(
                item.get("data_recebimento") or item.get("data")
            ) or date.today()

            if not nome_doador:
                continue

            doador, _ = Doador.objects.get_or_create(
                nome=nome_doador,
                defaults={
                    "tipo": "PF",
                    "telefone": normalizar_texto(item.get("telefone")),
                    "email": normalizar_texto(item.get("email")),
                    "cidade": normalizar_texto(item.get("cidade")),
                    "endereco": normalizar_texto(item.get("endereco")),
                }
            )

            campanha = None

            if nome_campanha:
                campanha, _ = Campanha.objects.get_or_create(
                    nome=nome_campanha,
                    defaults={
                        "descricao": "",
                        "data_inicio": data_recebimento,
                        "data_fim": data_recebimento,
                        "ativa": True,
                    }
                )

            Doacao.objects.create(
                doador=doador,
                campanha=campanha,
                data_recebimento=data_recebimento,
                observacao=normalizar_texto(item.get("observacao")),
                criado_por=usuario,
            )

            total += 1

    return total


# =========================
# DASHBOARD
# =========================

@login_required
def dashboard(request):
    doacoes_por_mes = (
        Doacao.objects
        .annotate(mes=TruncMonth("data_recebimento"))
        .values("mes")
        .annotate(total=Count("id"))
        .order_by("mes")
    )

    campanhas = (
        Campanha.objects
        .annotate(total_doacoes=Count("doacao"))
        .order_by("-total_doacoes")[:5]
    )

    produtos_estoque = Produto.objects.all()

    produtos_baixo_estoque = [
        produto for produto in produtos_estoque
        if produto.estoque_atual() <= produto.estoque_minimo
    ]

    context = {
        "total_doadores": Doador.objects.count(),
        "total_produtos": Produto.objects.count(),
        "total_doacoes": Doacao.objects.count(),
        "total_campanhas": Campanha.objects.filter(ativa=True).count(),
        "total_beneficiarios": Beneficiario.objects.count(),
        "produtos_baixo_estoque": produtos_baixo_estoque,
        "ultimas_doacoes": Doacao.objects.order_by("-criado_em")[:5],
        "grafico_meses": json.dumps([
            item["mes"].strftime("%m/%Y") for item in doacoes_por_mes
        ]),
        "grafico_doacoes": json.dumps([
            item["total"] for item in doacoes_por_mes
        ]),
        "grafico_campanhas_labels": json.dumps([
            item.nome for item in campanhas
        ]),
        "grafico_campanhas_valores": json.dumps([
            item.total_doacoes for item in campanhas
        ]),
    }

    return render(request, "gestao/dashboard.html", context)


# =========================
# CRUD GENÉRICO
# =========================

def crud_list(request, model, template, context_name):
    busca = request.GET.get("busca")

    if busca:
        dados = model.objects.filter(nome__icontains=busca)
    else:
        dados = model.objects.all()

    return render(request, template, {context_name: dados})


@login_required
def doadores(request):
    return crud_list(request, Doador, "gestao/doadores.html", "dados")


@login_required
def produtos(request):
    return crud_list(request, Produto, "gestao/produtos.html", "dados")


@login_required
def campanhas(request):
    return crud_list(request, Campanha, "gestao/campanhas.html", "dados")


@login_required
def doacoes(request):
    dados = Doacao.objects.all().order_by("-data_recebimento")
    return render(request, "gestao/doacoes.html", {"dados": dados})


# =========================
# FORMULÁRIOS
# =========================

@login_required
def criar_doador(request):
    form = DoadorForm(request.POST or None)

    if form.is_valid():
        form.save()
        messages.success(request, "Doador cadastrado com sucesso.")
        return redirect("doadores")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Novo Doador"
    })


@login_required
def editar_doador(request, pk):
    item = get_object_or_404(Doador, pk=pk)
    form = DoadorForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Doador atualizado com sucesso.")
        return redirect("doadores")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Doador"
    })


@login_required
def excluir_doador(request, pk):
    item = get_object_or_404(Doador, pk=pk)
    item.delete()
    messages.success(request, "Doador excluído com sucesso.")
    return redirect("doadores")


@login_required
def criar_produto(request):
    form = ProdutoForm(request.POST or None)

    if form.is_valid():
        form.save()
        messages.success(request, "Produto cadastrado com sucesso.")
        return redirect("produtos")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Novo Produto"
    })


@login_required
def editar_produto(request, pk):
    item = get_object_or_404(Produto, pk=pk)
    form = ProdutoForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Produto atualizado com sucesso.")
        return redirect("produtos")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Produto"
    })


@login_required
def excluir_produto(request, pk):
    item = get_object_or_404(Produto, pk=pk)
    item.delete()
    messages.success(request, "Produto excluído com sucesso.")
    return redirect("produtos")


@login_required
def criar_campanha(request):
    form = CampanhaForm(request.POST or None)

    if form.is_valid():
        form.save()
        messages.success(request, "Campanha cadastrada com sucesso.")
        return redirect("campanhas")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Nova Campanha"
    })


@login_required
def editar_campanha(request, pk):
    item = get_object_or_404(Campanha, pk=pk)
    form = CampanhaForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Campanha atualizada com sucesso.")
        return redirect("campanhas")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Campanha"
    })


@login_required
def excluir_campanha(request, pk):
    item = get_object_or_404(Campanha, pk=pk)
    item.delete()
    messages.success(request, "Campanha excluída com sucesso.")
    return redirect("campanhas")


@login_required
def criar_doacao(request):
    form = DoacaoForm(request.POST or None)

    if form.is_valid():
        doacao = form.save(commit=False)
        doacao.criado_por = request.user
        doacao.save()

        messages.success(request, "Doação cadastrada com sucesso.")
        return redirect("doacoes")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Nova Doação"
    })


@login_required
def editar_doacao(request, pk):
    item = get_object_or_404(Doacao, pk=pk)
    form = DoacaoForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Doação atualizada com sucesso.")
        return redirect("doacoes")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Doação"
    })


@login_required
def excluir_doacao(request, pk):
    item = get_object_or_404(Doacao, pk=pk)
    item.delete()
    messages.success(request, "Doação excluída com sucesso.")
    return redirect("doacoes")


# =========================
# RELATÓRIOS
# =========================

def get_model_data(tipo, request=None):
    data_inicio = request.GET.get("data_inicio") if request else None
    data_fim = request.GET.get("data_fim") if request else None
    campanha_id = request.GET.get("campanha") if request else None
    produto_id = request.GET.get("produto") if request else None
    doador_id = request.GET.get("doador") if request else None

    if tipo == "doadores":
        dados = Doador.objects.all()
        return dados, ["Nome", "Tipo", "Telefone", "Email", "Cidade"]

    if tipo == "produtos":
        dados = Produto.objects.all()
        return dados, ["Nome", "Categoria", "Unidade", "Estoque mínimo"]

    if tipo == "campanhas":
        dados = Campanha.objects.all()

        if data_inicio:
            dados = dados.filter(data_inicio__gte=data_inicio)

        if data_fim:
            dados = dados.filter(data_fim__lte=data_fim)

        return dados, ["Nome", "Início", "Fim", "Ativa"]

    if tipo == "doacoes":
        dados = Doacao.objects.all()

        if data_inicio:
            dados = dados.filter(data_recebimento__gte=data_inicio)

        if data_fim:
            dados = dados.filter(data_recebimento__lte=data_fim)

        if campanha_id:
            dados = dados.filter(campanha_id=campanha_id)

        if doador_id:
            dados = dados.filter(doador_id=doador_id)

        return dados, ["Doador", "Campanha", "Data", "Observação"]

    if tipo == "beneficiarios":
        dados = Beneficiario.objects.all()
        return dados, ["Nome", "Telefone", "Pessoas", "Ativo"]

    if tipo == "coletas":
        dados = Coleta.objects.all()

        if data_inicio:
            dados = dados.filter(data_agendada__date__gte=data_inicio)

        if data_fim:
            dados = dados.filter(data_agendada__date__lte=data_fim)

        if doador_id:
            dados = dados.filter(doador_id=doador_id)

        return dados, ["Doador", "Data", "Status", "Endereço"]

    if tipo == "estoque":
        dados = MovimentacaoEstoque.objects.all()

        if data_inicio:
            dados = dados.filter(data__date__gte=data_inicio)

        if data_fim:
            dados = dados.filter(data__date__lte=data_fim)

        if produto_id:
            dados = dados.filter(produto_id=produto_id)

        return dados, ["Produto", "Tipo", "Quantidade", "Data"]

    return [], []


def linha_objeto(tipo, obj):
    if tipo == "doadores":
        return [
            obj.nome,
            obj.get_tipo_display(),
            obj.telefone,
            obj.email,
            obj.cidade,
        ]

    if tipo == "produtos":
        return [
            obj.nome,
            obj.categoria.nome,
            obj.get_unidade_display(),
            obj.estoque_minimo,
        ]

    if tipo == "campanhas":
        return [
            obj.nome,
            obj.data_inicio,
            obj.data_fim,
            "Sim" if obj.ativa else "Não",
        ]

    if tipo == "doacoes":
        return [
            obj.doador.nome,
            obj.campanha.nome if obj.campanha else "Sem campanha",
            obj.data_recebimento,
            obj.observacao,
        ]

    if tipo == "beneficiarios":
        return [
            obj.nome,
            obj.telefone,
            obj.quantidade_pessoas,
            "Sim" if obj.ativo else "Não",
        ]

    if tipo == "coletas":
        return [
            obj.doador.nome,
            obj.data_agendada,
            obj.get_status_display(),
            obj.endereco_coleta,
        ]

    if tipo == "estoque":
        return [
            obj.produto.nome,
            obj.get_tipo_display(),
            obj.quantidade,
            obj.data,
        ]

    return []


@login_required
def exportar_excel(request, tipo):
    dados, colunas = get_model_data(tipo, request)

    wb = Workbook()
    ws = wb.active
    ws.title = f"Relatório {tipo}"

    ws.append(colunas)

    for obj in dados:
        ws.append(linha_objeto(tipo, obj))

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    response["Content-Disposition"] = f'attachment; filename="relatorio_{tipo}.xlsx"'

    wb.save(response)
    return response


@login_required
def exportar_word(request, tipo):
    dados, colunas = get_model_data(tipo, request)

    doc = Document()
    doc.add_heading(f"Relatório de {tipo.title()}", 0)

    table = doc.add_table(rows=1, cols=len(colunas))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells

    for i, coluna in enumerate(colunas):
        hdr_cells[i].text = coluna

    for obj in dados:
        row_cells = table.add_row().cells
        linha = linha_objeto(tipo, obj)

        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    response["Content-Disposition"] = f'attachment; filename="relatorio_{tipo}.docx"'

    doc.save(response)
    return response


@login_required
def exportar_pdf(request, tipo):
    dados, colunas = get_model_data(tipo, request)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="relatorio_{tipo}.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, 800, f"Relatório de {tipo.title()}")

    y = 760
    p.setFont("Helvetica-Bold", 9)

    x = 50
    for coluna in colunas:
        p.drawString(x, y, str(coluna))
        x += 120

    y -= 20
    p.setFont("Helvetica", 8)

    for obj in dados:
        x = 50
        linha = linha_objeto(tipo, obj)

        for valor in linha:
            p.drawString(x, y, str(valor)[:18])
            x += 120

        y -= 20

        if y < 50:
            p.showPage()
            y = 800

    p.save()
    return response


# =========================
# ESTOQUE
# =========================

@login_required
def estoque(request):
    dados = MovimentacaoEstoque.objects.all().order_by("-data")
    produtos = Produto.objects.all().order_by("nome")

    return render(request, "gestao/estoque.html", {
        "dados": dados,
        "produtos": produtos,
    })


@login_required
def criar_movimentacao(request):
    form = MovimentacaoEstoqueForm(request.POST or None)

    if form.is_valid():
        movimentacao = form.save(commit=False)

        if movimentacao.tipo == "SAIDA":
            if not movimentacao.produto.tem_estoque_suficiente(movimentacao.quantidade):
                messages.error(
                    request,
                    f"Estoque insuficiente. Estoque atual: {movimentacao.produto.estoque_atual()}"
                )
                return render(request, "gestao/form.html", {
                    "form": form,
                    "titulo": "Nova Movimentação de Estoque"
                })

        movimentacao.save()
        messages.success(request, "Movimentação registrada com sucesso.")
        return redirect("estoque")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Nova Movimentação de Estoque"
    })


@login_required
def excluir_movimentacao(request, pk):
    item = get_object_or_404(MovimentacaoEstoque, pk=pk)
    item.delete()
    messages.success(request, "Movimentação excluída com sucesso.")
    return redirect("estoque")


# =========================
# BENEFICIÁRIOS
# =========================

@login_required
def beneficiarios(request):
    busca = request.GET.get("busca")

    if busca:
        dados = Beneficiario.objects.filter(nome__icontains=busca)
    else:
        dados = Beneficiario.objects.all().order_by("nome")

    return render(request, "gestao/beneficiarios.html", {"dados": dados})


@login_required
def criar_beneficiario(request):
    form = BeneficiarioForm(request.POST or None)

    if form.is_valid():
        form.save()
        messages.success(request, "Beneficiário cadastrado com sucesso.")
        return redirect("beneficiarios")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Novo Beneficiário"
    })


@login_required
def editar_beneficiario(request, pk):
    item = get_object_or_404(Beneficiario, pk=pk)
    form = BeneficiarioForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Beneficiário atualizado com sucesso.")
        return redirect("beneficiarios")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Beneficiário"
    })


@login_required
def excluir_beneficiario(request, pk):
    item = get_object_or_404(Beneficiario, pk=pk)
    item.delete()
    messages.success(request, "Beneficiário excluído com sucesso.")
    return redirect("beneficiarios")


# =========================
# COLETAS
# =========================

@login_required
def coletas(request):
    dados = Coleta.objects.all().order_by("-data_agendada")
    return render(request, "gestao/coletas.html", {"dados": dados})


@login_required
def criar_coleta(request):
    form = ColetaForm(request.POST or None)

    if form.is_valid():
        form.save()
        messages.success(request, "Coleta cadastrada com sucesso.")
        return redirect("coletas")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Nova Coleta"
    })


@login_required
def editar_coleta(request, pk):
    item = get_object_or_404(Coleta, pk=pk)
    form = ColetaForm(request.POST or None, instance=item)

    if form.is_valid():
        form.save()
        messages.success(request, "Coleta atualizada com sucesso.")
        return redirect("coletas")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Editar Coleta"
    })


@login_required
def excluir_coleta(request, pk):
    item = get_object_or_404(Coleta, pk=pk)
    item.delete()
    messages.success(request, "Coleta excluída com sucesso.")
    return redirect("coletas")


# =========================
# ENTREGAS
# =========================

@login_required
def entregas(request):
    dados = EntregaBeneficiario.objects.all().order_by("-data_entrega")
    return render(request, "gestao/entregas.html", {"dados": dados})


@login_required
def criar_entrega(request):
    form = EntregaBeneficiarioForm(request.POST or None)

    if form.is_valid():
        entrega = form.save(commit=False)

        if not entrega.produto.tem_estoque_suficiente(entrega.quantidade):
            messages.error(
                request,
                f"Estoque insuficiente para entregar {entrega.produto.nome}. Estoque atual: {entrega.produto.estoque_atual()}"
            )
            return render(request, "gestao/form.html", {
                "form": form,
                "titulo": "Nova Entrega para Beneficiário"
            })

        entrega.save()

        MovimentacaoEstoque.objects.create(
            produto=entrega.produto,
            tipo="SAIDA",
            quantidade=entrega.quantidade,
            observacao=f"Entrega para beneficiário: {entrega.beneficiario.nome}"
        )

        messages.success(request, "Entrega cadastrada e estoque atualizado.")
        return redirect("entregas")

    return render(request, "gestao/form.html", {
        "form": form,
        "titulo": "Nova Entrega para Beneficiário"
    })


@login_required
def excluir_entrega(request, pk):
    item = get_object_or_404(EntregaBeneficiario, pk=pk)
    item.delete()
    messages.success(request, "Entrega excluída com sucesso.")
    return redirect("entregas")


# =========================
# RELATÓRIOS
# =========================

@login_required
def relatorios(request):
    context = {
        "doadores": Doador.objects.all().order_by("nome"),
        "campanhas": Campanha.objects.all().order_by("nome"),
        "produtos": Produto.objects.all().order_by("nome"),
    }

    return render(request, "gestao/relatorios.html", context)


# =========================
# DOCUMENTOS / IMPORTAÇÃO
# =========================

def extrair_texto_excel(caminho):
    workbook = openpyxl.load_workbook(caminho, data_only=True)
    texto = []

    for aba in workbook.sheetnames:
        sheet = workbook[aba]
        texto.append(f"Aba: {aba}")

        for linha in sheet.iter_rows(values_only=True):
            valores = [str(celula) for celula in linha if celula is not None]

            if valores:
                texto.append(" | ".join(valores))

    return "\n".join(texto)


def extrair_texto_word(caminho):
    documento = Document(caminho)
    texto = []

    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip():
            texto.append(paragrafo.text.strip())

    return "\n".join(texto)


def extrair_texto_pdf(caminho):
    reader = PdfReader(caminho)
    texto = []

    for pagina in reader.pages:
        conteudo = pagina.extract_text()

        if conteudo:
            texto.append(conteudo)

    return "\n".join(texto)


@login_required
def documentos(request):
    documentos = DocumentoImportado.objects.all().order_by("-importado_em")

    return render(request, "gestao/documentos.html", {
        "documentos": documentos
    })


@login_required
def importar_documentos(request):
    if request.method == "POST":
        form = DocumentoImportadoForm(request.POST, request.FILES)

        if form.is_valid():
            documento = form.save(commit=False)

            nome_arquivo = documento.arquivo.name.lower()

            if nome_arquivo.endswith(".xlsx"):
                documento.tipo = "EXCEL"
            elif nome_arquivo.endswith(".docx"):
                documento.tipo = "WORD"
            elif nome_arquivo.endswith(".pdf"):
                documento.tipo = "PDF"
            else:
                documento.tipo = "OUTRO"

            documento.save()
            caminho = documento.arquivo.path

            try:
                total_importado = 0

                if documento.tipo == "EXCEL":
                    texto = extrair_texto_excel(caminho)

                    if documento.tipo_importacao == "PRODUTOS":
                        total_importado = importar_produtos_excel(caminho)

                    elif documento.tipo_importacao == "DOADORES":
                        total_importado = importar_doadores_excel(caminho)

                    elif documento.tipo_importacao == "BENEFICIARIOS":
                        total_importado = importar_beneficiarios_excel(caminho)

                    elif documento.tipo_importacao == "CAMPANHAS":
                        total_importado = importar_campanhas_excel(caminho)

                    elif documento.tipo_importacao == "ESTOQUE":
                        total_importado = importar_estoque_excel(caminho)

                    elif documento.tipo_importacao == "DOACOES":
                        total_importado = importar_doacoes_excel(caminho, request.user)

                elif documento.tipo == "WORD":
                    texto = extrair_texto_word(caminho)

                elif documento.tipo == "PDF":
                    texto = extrair_texto_pdf(caminho)

                else:
                    texto = ""

                documento.texto_extraido = texto
                documento.save()

                RegistroImportado.objects.filter(documento=documento).delete()

                for linha in texto.splitlines():
                    if linha.strip():
                        RegistroImportado.objects.create(
                            documento=documento,
                            conteudo=linha.strip()
                        )

                if documento.tipo == "EXCEL":
                    messages.success(
                        request,
                        f"Documento importado com sucesso. {total_importado} registro(s) alimentaram o sistema."
                    )
                else:
                    messages.success(
                        request,
                        "Documento anexado e texto extraído. Para alimentar cadastros automaticamente, use planilha Excel."
                    )

                return redirect("importar_documentos")

            except Exception as erro:
                messages.error(request, f"Erro ao processar documento: {erro}")

    else:
        form = DocumentoImportadoForm()

    documentos = DocumentoImportado.objects.all().order_by("-importado_em")

    return render(request, "gestao/importar_documentos.html", {
        "form": form,
        "documentos": documentos
    })